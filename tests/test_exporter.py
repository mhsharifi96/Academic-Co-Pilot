"""
Unit tests for the .docx assembly behind compile_paper
(app/tools/exporter._build_docx). Pure filesystem — no network/LLM/DB.

Fixture-free so the suite also runs under tests/run_all.py.
"""

import os
import tempfile

from app.tools.exporter import _build_docx, _slugify


def test_slugify():
    assert _slugify("Agentic RAG for Legal Compliance!") == "agentic_rag_for_legal_compliance"
    assert _slugify("   ") == "paper"
    assert _slugify("") == "paper"


def test_build_docx_writes_title_and_sections():
    from docx import Document

    d = tempfile.mkdtemp()
    path = os.path.join(d, "out.docx")
    sections = [
        {"heading": "Introduction", "body": "First para.\n\nSecond para."},
        {"heading": "Conclusion", "body": "Wrap up."},
    ]
    _build_docx("My Paper", sections, path)
    assert os.path.isfile(path)

    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert "My Paper" in texts
    assert "Introduction" in texts
    assert "Conclusion" in texts
    assert "First para." in texts
    assert "Second para." in texts   # double-newline split into two paragraphs
    assert "Wrap up." in texts


def test_build_docx_creates_missing_parent_dir():
    from docx import Document

    d = tempfile.mkdtemp()
    path = os.path.join(d, "nested", "deeper", "out.docx")
    _build_docx("T", [{"heading": "H", "body": "B"}], path)
    assert os.path.isfile(path)
    assert "H" in [p.text for p in Document(path).paragraphs]
