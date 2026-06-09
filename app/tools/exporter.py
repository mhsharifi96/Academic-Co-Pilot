"""
Document export / assembly.

``compile_paper`` stitches drafted sections into a single Word (.docx) document
written under ``data/`` (served read-only by the existing ``/download`` endpoint
— see ``app/api/v1/endpoints/files.py``).

It writes to disk, so it IS gated by HITL (registered in ``INTERRUPT_TOOLS``).
"""

import os
import re
from typing import Dict, List, Optional

from langchain_core.tools import tool

# Match the upload/output directory the download endpoint already exposes.
OUTPUT_DIR = "data"


def _slugify(title: str) -> str:
    """Turn a paper title into a safe filename stem."""
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", (title or "paper").strip()).strip("_").lower()
    return slug or "paper"


def _build_docx(title: str, sections: List[Dict[str, str]], path: str) -> str:
    """
    Write a titled .docx with one heading + body per section. Pure filesystem
    work (no network/LLM) so it can be unit-tested offline.
    """
    # Imported lazily so app startup doesn't hard-depend on python-docx being
    # installed unless this tool is actually used.
    from docx import Document

    doc = Document()
    doc.add_heading(title or "Untitled", level=0)
    for section in sections:
        heading = (section.get("heading") or "").strip()
        body = (section.get("body") or "").strip()
        if heading:
            doc.add_heading(heading, level=1)
        if body:
            for para in body.split("\n\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc.save(path)
    return path


@tool
def compile_paper(
    title: str,
    sections: List[Dict[str, str]],
    output_path: Optional[str] = None,
) -> str:
    """
    Assemble drafted sections into a single Word (.docx) document for download.

    Use this AFTER the user has approved the individual sections (e.g. via the
    full-paper drafting flow) to produce the final deliverable.

    Args:
        title: The paper title (used as the document heading and filename).
        sections: Ordered list of sections, each a dict with 'heading' and
            'body' keys, e.g. [{"heading": "Introduction", "body": "..."}].
        output_path: Optional .docx path. Defaults to 'data/<slug>.docx'.
    """
    if not sections:
        return "Error: no sections provided to compile."

    if not output_path:
        output_path = os.path.join(OUTPUT_DIR, f"{_slugify(title)}.docx")
    elif not output_path.lower().endswith(".docx"):
        output_path += ".docx"

    try:
        path = _build_docx(title, sections, output_path)
    except Exception as e:
        return f"Error compiling document: {str(e)}"

    return (
        f"Compiled {len(sections)} sections into '{path}'. "
        f"It can be downloaded via the /download endpoint."
    )
